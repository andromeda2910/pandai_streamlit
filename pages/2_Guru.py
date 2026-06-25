import streamlit as st
from google import genai
from google.genai import types
import os
import time
import json
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# =================================================
# 1. KONFIGURASI HALAMAN
# =================================================
st.set_page_config(
    page_title="PANDAI - Guru",
    page_icon="👨‍🏫",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Load CSS
def local_css(file_name):
    if os.path.exists(file_name):
        with open(file_name, "r") as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

local_css("style.css")

# Custom Loading Component
def show_loading(message="Harap tunggu..."):
    loading_html = f"""
    <div style='
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        padding: 40px;
        background: rgba(255, 255, 255, 0.95);
        border-radius: 16px;
        border: 2px solid #e5e7eb;
        margin: 20px 0;
    '>
        <div style='
            width: 60px;
            height: 60px;
            border: 5px solid #e5e7eb;
            border-top: 5px solid #3b82f6;
            border-radius: 50%;
            animation: spin 1s linear infinite;
            margin-bottom: 20px;
        '></div>
        <p style='
            font-size: 24px;
            font-weight: 600;
            color: #1f2937;
            margin: 0;
            animation: pulse 1.5s ease-in-out infinite;
        '>{message}</p>
        <style>
            @keyframes spin {{
                0% {{ transform: rotate(0deg); }}
                100% {{ transform: rotate(360deg); }}
            }}
            @keyframes pulse {{
                0%, 100% {{ opacity: 1; }}
                50% {{ opacity: 0.5; }}
            }}
        </style>
    </div>
    """
    return loading_html

# =================================================
# 2. HEADER & HOME BUTTON
# =================================================

# Home button at top
col_home1, col_home2, col_home3 = st.columns([1, 2, 1])
with col_home2:
    if st.button("🏠 Kembali ke Beranda", use_container_width=True, key="guru_home"):
        st.switch_page("app.py")

st.markdown('<p class="hero-title"><span class="hero-icon">👨‍🏫</span>Guru</p>', unsafe_allow_html=True)
st.markdown('<div class="custom-divider"></div>', unsafe_allow_html=True)

# SD-only notice
st.markdown("""
    <div style='background: #FEF3C7; padding: 15px; border-radius: 12px; border-left: 4px solid #F59E0B; margin: 20px 0;'>
        <p style='color: #92400E; font-weight: 600; margin: 0;'>📚 Perhatian: Saat ini materi tersedia untuk jenjang SD (Kelas 1-6). Jenjang lain menyusul.</p>
    </div>
""", unsafe_allow_html=True)

# =================================================
# 3. INITIALIZE NAVIGATION STATE
# =================================================
if "menu_aktif" not in st.session_state:
    st.session_state.menu_aktif = "rpp"

# =================================================
# 4. MENU NAVIGASI TOMBOL (MODERN ESTATIK)
# =================================================
st.markdown("### 🛠️ Pilih Fitur Yang Ingin Digunakan")
col_menu1, col_menu2 = st.columns(2)

with col_menu1:
    tipe_rpp = "primary" if st.session_state.menu_aktif == "rpp" else "secondary"
    if st.button("📝 Perencanaan Pembelajaran (RPP)", use_container_width=True, type=tipe_rpp, key="btn_nav_rpp"):
        st.session_state.menu_aktif = "rpp"
        st.rerun()

with col_menu2:
    tipe_kuis = "primary" if st.session_state.menu_aktif == "kuis" else "secondary"
    if st.button("🎯 Generator Kuis (LKPD)", use_container_width=True, type=tipe_kuis, key="btn_nav_kuis"):
        st.session_state.menu_aktif = "kuis"
        st.rerun()

st.divider()

# API Key Validation (Global check for both features)
api_key = os.getenv("GEMINI_API_KEY")

# =================================================
# 5. KONDISI TAMPILAN BERDASARKAN TOMBOL AKTIF
# =================================================

# -------------------------------------------------
# MENU A: GENERATOR RPP
# -------------------------------------------------
if st.session_state.menu_aktif == "rpp":
    st.markdown("## 📝 Perencanaan Pembelajaran (RPP)")
    st.caption("💡 Pembuatan Rencana Pelaksanaan Pembelajaran otomatis berbasis AI sesuai standar Kurikulum Nasional.")
    
    # Form Input RPP
    col1, col2 = st.columns([1, 1])
    with col1:
        mata_pelajaran = st.selectbox(
            "📚 Mata Pelajaran",
            ["Matematika", "IPA", "Bahasa Indonesia", "IPS", "PKn", "Bahasa Inggris", "Seni Budaya", "PJOK"],
            index=0
        )
        kelas = st.selectbox(
            "🎓 Kelas",
            ["Kelas 1 SD", "Kelas 2 SD", "Kelas 3 SD", "Kelas 4 SD", "Kelas 5 SD", "Kelas 6 SD"],
            index=3
        )
    with col2:
        topik = st.text_input(
            "📌 Topik Pembelajaran",
            placeholder="Contoh: Penjumlahan Pecahan, Siklus Air, dll."
        )
        durasi = st.selectbox(
            "⏱️ Durasi Pembelajaran",
            ["45 Menit (1 JP)", "60 Menit", "90 Menit (2 JP)", "120 Menit"],
            index=0
        )

    st.divider()

    # Generate Button
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    with col_btn2:
        generate_btn = st.button("🚀 Generate RPP", type="primary", use_container_width=True, key="generate_rpp")

    # Output Logic RPP
    if generate_btn:
        if not topik:
            st.error("❌ Mohon isi topik pembelajaran!")
        elif not api_key:
            st.warning("⚠️ GEMINI_API_KEY tidak ditemukan di .env file.")
        else:
            try:
                # Cooldown check
                if "last_rpp_request" in st.session_state:
                    time_since_last = time.time() - st.session_state.last_rpp_request
                    if time_since_last < 15:
                        remaining = int(15 - time_since_last)
                        st.error(f"❌ Mohon tunggu {remaining} detik sebelum membuat RPP lagi (untuk menghindari limit kuota)")
                        st.stop()
                
                client = genai.Client(api_key=api_key)
                loading_placeholder = st.empty()
                loading_placeholder.markdown(show_loading("Harap tunggu, AI sedang menyusun RPP..."), unsafe_allow_html=True)
                
                time.sleep(0.5)
                st.session_state.last_rpp_request = time.time()
                
                prompt = f"""
                    Kamu adalah seorang pakar kurikulum pendidikan di Indonesia. Buatlah Rencana Pelaksanaan Pembelajaran (RPP) formal tanpa kalimat pembuka, tanpa kata pengantar, tanpa basa-basi (seperti "Tentu, berikut adalah..."), dan tanpa kalimat penutup di akhir dokumen. 
                    
                    Langsung mulai dokumen dari teks berkepala: "RENCANA PELAKSANAAN PEMBELAJARAN (RPP)"
                    
                    Metadata Dokumen:
                    Mata Pelajaran: {mata_pelajaran}
                    Kelas: {kelas}
                    Topik: {topik}
                    Durasi: {durasi}
                    
                    Format RPP wajib mencakup komponen berikut secara berurutan:
                    1. TUJUAN PEMBELAJARAN (Tujuan umum dan khusus)
                    2. MATERI AJAR (Ringkasan materi yang akan diajarkan)
                    3. METODE PEMBELAJARAN (Metode yang digunakan dan langkah-langkahnya)
                    4. MEDIA DAN SUMBER BELAJAR (Alat dan bahan yang diperlukan)
                    5. LANGKAH-LANGKAH KEGIATAN (Pendahuluan, Inti, Penutup)
                    6. PENILAIAN (Teknik dan instrumen penilaian)
                    
                    Gunakan bahasa Indonesia yang baku, formal, instruksional, dan langsung siap pakai untuk administrasi sekolah.
                    """
                
                response = client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=prompt
                )
                
                loading_placeholder.empty()
                
                st.markdown("### 📄 RPP yang Dihasilkan")
                st.markdown(f"""
                    <div style='background: #FFFFFF; padding: 30px; border-radius: 16px; border: 2px solid #E5E7EB; margin: 20px 0;'>
                        <div style='white-space: pre-wrap; line-height: 1.8; color: #1F2937;'>{response.text}</div>
                    </div>
                """, unsafe_allow_html=True)
                
                st.success("✅ RPP berhasil dibuat!")
                
                # --- PROSES GENERASI FILE WORD (.DOCX) ---
                from docx import Document
                from docx.shared import Pt
                import io
                
                doc_io = io.BytesIO()
                doc = Document()
                
                # Set font global ke Arial (Standar Formal)
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Arial'
                font.size = Pt(11)
                
                # Judul RPP di dokumen Word
                title_p = doc.add_paragraph()
                title_run = title_p.add_run("RENCANA PELAKSANAAN PEMBELAJARAN (RPP)")
                title_run.bold = True
                title_run.font.size = Pt(14)
                title_p.alignment = 1  # Center alignment
                
                # Metadata RPP
                doc.add_paragraph(f"Mata Pelajaran: {mata_pelajaran}")
                doc.add_paragraph(f"Kelas / Jenjang: {kelas}")
                doc.add_paragraph(f"Topik / Materi: {topik}")
                doc.add_paragraph(f"Durasi Waktu: {durasi}")
                doc.add_paragraph("-" * 60)
                
                # Susun teks utama baris demi baris
                for line in response.text.split('\n'):
                    if line.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', 'A.', 'B.', 'C.')):
                        p = doc.add_paragraph()
                        r = p.add_run(line)
                        r.bold = True
                    else:
                        doc.add_paragraph(line)
                
                doc.save(doc_io)
                doc_bytes = doc_io.getvalue()
                
                # Tombol Download Word asli (.docx)
                st.download_button(
                    label="📥 Download RPP (Format Microsoft Word .docx)",
                    data=doc_bytes,
                    file_name=f"RPP_{mata_pelajaran}_{kelas}_{topik.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                # Append ke history RPP
                if "rpp_history" not in st.session_state:
                    st.session_state.rpp_history = []
                st.session_state.rpp_history.append({
                    "mapel": mata_pelajaran,
                    "kelas": kelas,
                    "topik": topik,
                    "content": response.text
                })
                    
            except Exception as e:
                loading_placeholder.empty()
                error_msg = str(e).lower()
                
                if "503" in error_msg or "unavailable" in error_msg:
                    st.error("⏳ Server PANDAI Sedang Sangat Padat")
                    st.info("""
                        💡 **Tips Penggunaan:** Saat ini banyak guru di seluruh Indonesia yang sedang menggunakan fitur ini secara bersamaan. 
                        Mohon tunggu sekitar **10-30 detik**, lalu tekan tombol **🚀 Generate RPP** kembali.
                    """)
                elif "quota" in error_msg or "429" in error_msg:
                    st.error("❌ Sistem Sedang Sibuk (Batasan Kuota)")
                    st.info("💡 Mohon tunggu beberapa saat sebelum mencoba membuat dokumen baru kembali.")
                else:
                    st.error("⚠️ Terjadi Kendala Teknis")
                    st.caption(f"Detail kendala: {e}")

    # Display History RPP
    if "rpp_history" in st.session_state and st.session_state.rpp_history:
        st.divider()
        st.markdown("## 📚 Riwayat RPP Berhasil")
        
        for i, rpp in enumerate(reversed(st.session_state.rpp_history), 1):
            with st.expander(f"RPP #{i}: {rpp['mapel']} - {rpp['kelas']} - {rpp['topik']}"):
                st.markdown(f"""
                    <div style='background: #F9FAFB; padding: 20px; border-radius: 12px; margin: 10px 0;'>
                        <p style='white-space: pre-wrap; line-height: 1.6; color: #1F2937;'>{rpp['content']}</p>
                    </div>
                """, unsafe_allow_html=True)
        
        st.write("") # Spacer
        if st.button("🗑️ Hapus Riwayat RPP", use_container_width=True, key="clear_rpp_history"):
            st.session_state.rpp_history = []
            st.rerun()

# -------------------------------------------------
# MENU B: GENERATOR KUIS (LKPD)
# -------------------------------------------------
else:
    st.markdown("## 📝 Generator Kuis (LKPD)")
    st.caption("💡 Buat kuis interaktif pilihan ganda dengan AI dan unduh langsung dalam format cetak PDF LKPD.")
    
    # Quiz Form
    col_q1, col_q2 = st.columns([1, 1])
    with col_q1:
        quiz_kelas = st.selectbox(
            "🎓 Kelas",
            ["Kelas 1 SD", "Kelas 2 SD", "Kelas 3 SD", "Kelas 4 SD", "Kelas 5 SD", "Kelas 6 SD"],
            index=3,
            key="quiz_kelas"
        )
        quiz_mapel = st.selectbox(
            "📚 Mata Pelajaran",
            ["Matematika", "IPA", "Bahasa Indonesia", "IPS", "PKn"],
            index=0,
            key="quiz_mapel"
        )
    with col_q2:
        quiz_jumlah = st.number_input(
            "📊 Jumlah Soal",
            min_value=1,
            max_value=20,
            value=5,
            step=1,
            key="quiz_jumlah"
        )
        quiz_kesulitan = st.selectbox(
            "🎯 Tingkat Kesulitan",
            ["Mudah", "Sedang", "Sulit"],
            index=1,
            key="quiz_kesulitan"
        )

    quiz_topik = st.text_input(
        "📌 Topik Pembelajaran",
        placeholder="Contoh: Penjumlahan, Siklus Air, dll.",
        key="quiz_topik"
    )

    st.divider()

    # Generate Quiz Button
    col_qbtn1, col_qbtn2, col_qbtn3 = st.columns([1, 2, 1])
    with col_qbtn2:
        generate_quiz_btn = st.button("🚀 Generate Kuis", type="primary", use_container_width=True, key="generate_quiz")

    # Output Logic Quiz
    if generate_quiz_btn:
        if not quiz_topik:
            st.error("❌ Mohon isi topik pembelajaran!")
        elif not api_key:
            st.error("❌ Ada kesalahan pada sistem API: KEY tidak terdaftar.")
        else:
            try:
                # Cooldown check
                if "last_quiz_request" in st.session_state:
                    time_since_last = time.time() - st.session_state.last_quiz_request
                    if time_since_last < 15:
                        remaining = int(15 - time_since_last)
                        st.error(f"❌ Mohon tunggu {remaining} detik sebelum membuat kuis lagi (untuk menghindari limit kuota)")
                        st.stop()
                
                client = genai.Client(api_key=api_key)
                loading_placeholder_quiz = st.empty()
                loading_placeholder_quiz.markdown(show_loading("Harap tunggu, AI sedang menyusun kuis interaktif..."), unsafe_allow_html=True)
                
                time.sleep(0.5)
                st.session_state.last_quiz_request = time.time()
                
                prompt = f"""
                    Bertindaklah sebagai pembuat soal ujian profesional. Buatlah tepat {quiz_jumlah} butir soal pilihan ganda (A, B, C, D) yang siap diujikan untuk mata pelajaran {quiz_mapel} kelas {quiz_kelas} mengenai topik "{quiz_topik}".
                    
                    Karakteristik Soal:
                    - Tingkat Kesulitan: {quiz_kesulitan}
                    - Bahasa: Menggunakan ragam bahasa baku Indonesia resmi, bersifat evaluatif, objektif, jelas, dan sesuai dengan indikator pencapaian kompetensi siswa {quiz_kelas}.
                    - Hindari segala bentuk teks pengantar di luar struktur data.
                    """
                
                response = client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        response_mime_type="application/json",
                        response_schema=types.Schema(
                            type=types.Type.OBJECT,
                            properties={
                                "soal": types.Schema(
                                    type=types.Type.ARRAY,
                                    items=types.Schema(
                                        type=types.Type.OBJECT,
                                        properties={
                                            "nomor": types.Schema(type=types.Type.INTEGER),
                                            "pertanyaan": types.Schema(type=types.Type.STRING),
                                            "pilihan": types.Schema(
                                                type=types.Type.OBJECT,
                                                properties={
                                                    "A": types.Schema(type=types.Type.STRING),
                                                    "B": types.Schema(type=types.Type.STRING),
                                                    "C": types.Schema(type=types.Type.STRING),
                                                    "D": types.Schema(type=types.Type.STRING),
                                                },
                                                required=["A", "B", "C", "D"]
                                            ),
                                            "jawaban_benar": types.Schema(type=types.Type.STRING)
                                        },
                                        required=["nomor", "pertanyaan", "pilihan", "jawaban_benar"]
                                    )
                                )
                            },
                            required=["soal"]
                        )
                    )
                )
                
                loading_placeholder_quiz.empty()
                
                try:
                    quiz_data = json.loads(response.text.strip())
                    st.success("✅ Kuis berhasil dibuat!")
                    
                    # Display quiz preview
                    st.markdown("### 📄 Preview Kuis")
                    for soal in quiz_data.get("soal", []):
                        st.markdown(f"""
                            <div style='background: #F9FAFB; padding: 20px; border-radius: 12px; margin: 15px 0; border-left: 4px solid #3B82F6;'>
                                <p style='font-weight: bold; color: #1F2937; margin: 0;'>Soal {soal['nomor']}</p>
                                <p style='color: #374151; margin: 10px 0;'>{soal['pertanyaan']}</p>
                                <div style='margin-left: 20px;'>
                                    <p style='color: #4B5563; margin: 5px 0;'>A. {soal['pilihan']['A']}</p>
                                    <p style='color: #4B5563; margin: 5px 0;'>B. {soal['pilihan']['B']}</p>
                                    <p style='color: #4B5563; margin: 5px 0;'>C. {soal['pilihan']['C']}</p>
                                    <p style='color: #4B5563; margin: 5px 0;'>D. {soal['pilihan']['D']}</p>
                                </div>
                                <p style='color: #059669; font-weight: 600; margin: 10px 0 0 0;'>✓ Jawaban: {soal['jawaban_benar']}</p>
                            </div>
                        """, unsafe_allow_html=True)
                    
                    # Generate PDF (ReportLab)
                    st.markdown("### 📥 Download Kuis sebagai PDF (LKPD)")
                    from reportlab.lib.pagesizes import A4
                    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet
                    from reportlab.lib.units import inch
                    import io
                    
                    buffer = io.BytesIO()
                    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=0.5*inch, rightMargin=0.5*inch, topMargin=0.5*inch, bottomMargin=0.5*inch)
                    story = []
                    styles = getSampleStyleSheet()
                    
                    story.append(Paragraph(f"<b>LEMBAR KERJA PESERTA DIDIK (LKPD)</b>", styles['Title']))
                    story.append(Spacer(1, 0.2*inch))
                    
                    info_data = [
                        ["Mata Pelajaran:", quiz_mapel],
                        ["Kelas:", quiz_kelas],
                        ["Topik:", quiz_topik],
                        ["Jumlah Soal:", str(quiz_jumlah)],
                        ["Tingkat Kesulitan:", quiz_kesulitan]
                    ]
                    info_table = Table(info_data, colWidths=[2*inch, 4*inch])
                    info_table.setStyle(TableStyle([
                        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 11),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ]))
                    story.append(info_table)
                    story.append(Spacer(1, 0.3*inch))
                    
                    for soal in quiz_data.get("soal", []):
                        story.append(Paragraph(f"<b>{soal['nomor']}. {soal['pertanyaan']}</b>", styles['BodyText']))
                        opts = [f"A. {soal['pilihan']['A']}", f"B. {soal['pilihan']['B']}", f"C. {soal['pilihan']['C']}", f"D. {soal['pilihan']['D']}"]
                        for opt in opts:
                            story.append(Paragraph(opt, styles['BodyText']))
                        story.append(Spacer(1, 0.15*inch))
                    
                    story.append(Spacer(1, 0.2*inch))
                    story.append(Paragraph("<b>KUNCI JAWABAN:</b>", styles['Heading3']))
                    answer_text = ""
                    for soal in quiz_data.get("soal", []):
                        answer_text += f"{soal['nomor']}. {soal['jawaban_benar']}  "
                    story.append(Paragraph(answer_text, styles['BodyText']))
                    
                    doc.build(story)
                    pdf_bytes = buffer.getvalue()
                    buffer.close()
                    
                    st.download_button(
                        label="📥 Download Kuis PDF",
                        data=pdf_bytes,
                        file_name=f"LKPD_{quiz_mapel}_{quiz_kelas}_{quiz_topik}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                    
                    if "quiz_history" not in st.session_state:
                        st.session_state.quiz_history = []
                    st.session_state.quiz_history.append({
                        "mapel": quiz_mapel,
                        "kelas": quiz_kelas,
                        "topik": quiz_topik,
                        "jumlah": quiz_jumlah,
                        "kesulitan": quiz_kesulitan,
                        "data": quiz_data
                    })
                    
                except json.JSONDecodeError:
                    st.error("❌ Gagal memformat kuis ke skema JSON. Silakan coba lagi.")
                    
            except Exception as e:
                loading_placeholder_quiz.empty()
                error_msg = str(e).lower()
                
                if "503" in error_msg or "unavailable" in error_msg:
                    st.error("⏳ Server PANDAI Sedang Sangat Padat")
                    st.info("""
                        💡 **Tips Penggunaan:** Banyak guru yang sedang membuat kuis saat ini. 
                        Mohon tunggu sekitar **10-30 detik**, lalu tekan tombol **🚀 Generate Kuis** kembali.
                    """)
                elif "quota" in error_msg or "429" in error_msg:
                    st.error("❌ Sistem Sedang Sibuk (Batasan Kuota)")
                    st.info("💡 Mohon tunggu beberapa saat sebelum mencoba membuat kuis kembali.")
                else:
                    st.error("⚠️ Terjadi Kendala Teknis")
                    st.caption(f"Detail kendala: {e}")

    # Display History Quiz
    if "quiz_history" in st.session_state and st.session_state.quiz_history:
        st.divider()
        st.markdown("## 📚 Riwayat Kuis Berhasil")
        
        for i, quiz in enumerate(reversed(st.session_state.quiz_history), 1):
            with st.expander(f"Kuis #{i}: {quiz['mapel']} - {quiz['kelas']} - {quiz['topik']} ({quiz['jumlah']} soal)"):
                st.markdown(f"""
                    <div style='background: #F9FAFB; padding: 15px; border-radius: 12px; margin: 10px 0;'>
                        <p><b>Mata Pelajaran:</b> {quiz['mapel']}</p>
                        <p><b>Kelas:</b> {quiz['kelas']}</p>
                        <p><b>Topik:</b> {quiz['topik']}</p>
                        <p><b>Jumlah Soal:</b> {quiz['jumlah']}</p>
                        <p><b>Tingkat Kesulitan:</b> {quiz['kesulitan']}</p>
                    </div>
                """, unsafe_allow_html=True)
                
        st.write("") # Spacer
        if st.button("🗑️ Hapus Riwayat Kuis", use_container_width=True, key="clear_quiz_history"):
            st.session_state.quiz_history = []
            st.rerun()

# =================================================
# 6. FOOTER
# =================================================
st.markdown("""
    <div style='text-align: center; color: #9CA3AF; font-size: 13px; padding-top: 25px; border-top: 1px solid #E5E7EB; margin-top: 40px;'>
        🐼 PANDAI - Mode Guru • © 2026
    </div>
""", unsafe_allow_html=True)